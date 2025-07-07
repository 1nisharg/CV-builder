import streamlit as st
import requests
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn
import io
import json
from datetime import datetime

# Configure Streamlit page
st.set_page_config(
    page_title="Professional CV Generator",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS for better styling
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #2E86AB;
        text-align: center;
        margin-bottom: 2rem;
    }
    .section-header {
        font-size: 1.5rem;
        font-weight: bold;
        color: #A23B72;
        margin-top: 2rem;
        margin-bottom: 1rem;
    }
    .info-box {
        background-color: #f0f2f6;
        padding: 1rem;
        border-radius: 0.5rem;
        margin: 1rem 0;
    }
    .success-box {
        background-color: #d4edda;
        color: #155724;
        padding: 1rem;
        border-radius: 0.5rem;
        border: 1px solid #c3e6cb;
    }
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'cv_data' not in st.session_state:
    st.session_state.cv_data = {}

def enhance_with_groq(text, section_type, groq_api_key):
    """Enhance text using Groq API for professional CV writing"""
    if not groq_api_key or not text.strip():
        return text
    
    prompts = {
        "profile": f"Enhance this professional profile/summary for a CV to make it more compelling and ATS-friendly while keeping it concise (2-3 sentences): {text}",
        "experience": f"Enhance this work experience description for a CV using action verbs, quantifiable achievements, and professional language: {text}",
        "education": f"Enhance this education information for a CV, making it more professional and comprehensive: {text}",
        "skills": f"Enhance and organize these skills for a CV, grouping them professionally and using industry-standard terminology: {text}",
        "activities": f"Enhance these activities and interests for a CV, focusing on professional relevance and leadership qualities: {text}"
    }
    
    try:
        headers = {
            "Authorization": f"Bearer {groq_api_key}",
            "Content-Type": "application/json"
        }
        
        data = {
            "messages": [
                {
                    "role": "user",
                    "content": prompts.get(section_type, f"Enhance this text for a professional CV: {text}")
                }
            ],
            "model": "llama3-8b-8192",
            "temperature": 0.7,
            "max_tokens": 500
        }
        
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers=headers,
            json=data,
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            return result['choices'][0]['message']['content'].strip()
        else:
            st.warning(f"Groq API error: {response.status_code}")
            return text
            
    except Exception as e:
        st.warning(f"Error enhancing text: {str(e)}")
        return text

def create_cv_document(cv_data, enhance_with_ai=False, groq_api_key=None):
    """Create a professional CV document using python-docx"""
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
    
    # Add name as header
    if cv_data.get('name'):
        name_paragraph = doc.add_paragraph()
        name_run = name_paragraph.add_run(cv_data['name'])
        name_run.font.size = Inches(0.3)
        name_run.bold = True
        name_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Add space
    
    # Add Profile section
    if cv_data.get('profile'):
        profile_heading = doc.add_heading('Profile', level=1)
        profile_text = cv_data['profile']
        if enhance_with_ai and groq_api_key:
            profile_text = enhance_with_groq(profile_text, "profile", groq_api_key)
        doc.add_paragraph(profile_text)
        doc.add_paragraph()
    
    # Add Experience section
    if cv_data.get('experience'):
        exp_heading = doc.add_heading('Experience', level=1)
        for i, exp in enumerate(cv_data['experience']):
            if exp.get('company') or exp.get('position') or exp.get('description'):
                # Job title and company
                if exp.get('position') and exp.get('company'):
                    job_para = doc.add_paragraph()
                    job_run = job_para.add_run(f"{exp['position']} - {exp['company']}")
                    job_run.bold = True
                
                # Duration
                if exp.get('duration'):
                    duration_para = doc.add_paragraph(exp['duration'])
                    duration_para.style = 'List Bullet'
                
                # Description
                if exp.get('description'):
                    desc_text = exp['description']
                    if enhance_with_ai and groq_api_key:
                        desc_text = enhance_with_groq(desc_text, "experience", groq_api_key)
                    doc.add_paragraph(desc_text)
                
                if i < len(cv_data['experience']) - 1:
                    doc.add_paragraph()  # Add space between experiences
        doc.add_paragraph()
    
    # Add Education section
    if cv_data.get('education'):
        edu_heading = doc.add_heading('Education', level=1)
        for i, edu in enumerate(cv_data['education']):
            if edu.get('degree') or edu.get('institution') or edu.get('details'):
                # Degree and institution
                if edu.get('degree') and edu.get('institution'):
                    edu_para = doc.add_paragraph()
                    edu_run = edu_para.add_run(f"{edu['degree']} - {edu['institution']}")
                    edu_run.bold = True
                
                # Year
                if edu.get('year'):
                    doc.add_paragraph(edu['year'])
                
                # Additional details
                if edu.get('details'):
                    details_text = edu['details']
                    if enhance_with_ai and groq_api_key:
                        details_text = enhance_with_groq(details_text, "education", groq_api_key)
                    doc.add_paragraph(details_text)
                
                if i < len(cv_data['education']) - 1:
                    doc.add_paragraph()  # Add space between education entries
        doc.add_paragraph()
    
    # Add Skills & Abilities section
    if cv_data.get('skills'):
        skills_heading = doc.add_heading('Skills & Abilities', level=1)
        skills_text = cv_data['skills']
        if enhance_with_ai and groq_api_key:
            skills_text = enhance_with_groq(skills_text, "skills", groq_api_key)
        doc.add_paragraph(skills_text)
        doc.add_paragraph()
    
    # Add Activities and Interests section
    if cv_data.get('activities'):
        activities_heading = doc.add_heading('Activities and Interests', level=1)
        activities_text = cv_data['activities']
        if enhance_with_ai and groq_api_key:
            activities_text = enhance_with_groq(activities_text, "activities", groq_api_key)
        doc.add_paragraph(activities_text)
    
    return doc

def main():
    st.markdown('<div class="main-header">🚀 Professional CV Generator</div>', unsafe_allow_html=True)
    st.markdown("Transform your information into a professional CV optimized for top companies")
    
    # Sidebar for API configuration
    with st.sidebar:
        st.header("⚙️ Configuration")
        groq_api_key = st.text_input("Groq API Key (Optional)", type="password", 
                                   help="Enter your Groq API key for AI-enhanced CV writing")
        enhance_with_ai = st.checkbox("Enhance with AI", value=False, 
                                    help="Use AI to improve your CV content")
        
        if enhance_with_ai and not groq_api_key:
            st.warning("Please enter your Groq API key to use AI enhancement")
        
        st.markdown("---")
        st.markdown("### 📋 How to use:")
        st.markdown("1. Fill in your information")
        st.markdown("2. Optionally enable AI enhancement")
        st.markdown("3. Generate your professional CV")
        st.markdown("4. Download as Word document")
    
    # Main content area
    col1, col2 = st.columns([2, 1])
    
    with col1:
        # Personal Information
        st.markdown('<div class="section-header">👤 Personal Information</div>', unsafe_allow_html=True)
        name = st.text_input("Full Name *", placeholder="John Smith")
        
        # Profile Section
        st.markdown('<div class="section-header">📝 Profile</div>', unsafe_allow_html=True)
        profile = st.text_area("Professional Profile/Summary *", 
                             placeholder="Brief professional summary highlighting your key strengths and career objectives...",
                             height=100)
        
        # Experience Section
        st.markdown('<div class="section-header">💼 Experience</div>', unsafe_allow_html=True)
        
        # Dynamic experience entries
        if 'num_experiences' not in st.session_state:
            st.session_state.num_experiences = 1
        
        experiences = []
        for i in range(st.session_state.num_experiences):
            with st.expander(f"Experience {i+1}", expanded=True):
                col_exp1, col_exp2 = st.columns(2)
                with col_exp1:
                    position = st.text_input(f"Position", key=f"pos_{i}", placeholder="Software Engineer")
                    company = st.text_input(f"Company", key=f"comp_{i}", placeholder="Tech Corp")
                with col_exp2:
                    duration = st.text_input(f"Duration", key=f"dur_{i}", placeholder="Jan 2020 - Present")
                
                description = st.text_area(f"Description", key=f"desc_{i}",
                                         placeholder="Key responsibilities and achievements...",
                                         height=80)
                
                experiences.append({
                    'position': position,
                    'company': company,
                    'duration': duration,
                    'description': description
                })
        
        col_exp_btn1, col_exp_btn2 = st.columns(2)
        with col_exp_btn1:
            if st.button("➕ Add Experience"):
                st.session_state.num_experiences += 1
                st.experimental_rerun()
        with col_exp_btn2:
            if st.button("➖ Remove Experience") and st.session_state.num_experiences > 1:
                st.session_state.num_experiences -= 1
                st.experimental_rerun()
        
        # Education Section
        st.markdown('<div class="section-header">🎓 Education</div>', unsafe_allow_html=True)
        
        # Dynamic education entries
        if 'num_education' not in st.session_state:
            st.session_state.num_education = 1
        
        education = []
        for i in range(st.session_state.num_education):
            with st.expander(f"Education {i+1}", expanded=True):
                col_edu1, col_edu2 = st.columns(2)
                with col_edu1:
                    degree = st.text_input(f"Degree", key=f"degree_{i}", placeholder="Bachelor of Science")
                    institution = st.text_input(f"Institution", key=f"inst_{i}", placeholder="University Name")
                with col_edu2:
                    year = st.text_input(f"Year", key=f"year_{i}", placeholder="2020")
                
                details = st.text_area(f"Additional Details", key=f"edu_details_{i}",
                                     placeholder="GPA, relevant coursework, honors...",
                                     height=60)
                
                education.append({
                    'degree': degree,
                    'institution': institution,
                    'year': year,
                    'details': details
                })
        
        col_edu_btn1, col_edu_btn2 = st.columns(2)
        with col_edu_btn1:
            if st.button("➕ Add Education"):
                st.session_state.num_education += 1
                st.experimental_rerun()
        with col_edu_btn2:
            if st.button("➖ Remove Education") and st.session_state.num_education > 1:
                st.session_state.num_education -= 1
                st.experimental_rerun()
        
        # Skills Section
        st.markdown('<div class="section-header">🛠️ Skills & Abilities</div>', unsafe_allow_html=True)
        skills = st.text_area("Skills & Abilities *", 
                            placeholder="Technical skills, programming languages, tools, soft skills...",
                            height=100)
        
        # Activities Section
        st.markdown('<div class="section-header">🏆 Activities and Interests</div>', unsafe_allow_html=True)
        activities = st.text_area("Activities and Interests", 
                                placeholder="Professional activities, hobbies, volunteer work...",
                                height=80)
    
    with col2:
        st.markdown('<div class="section-header">📄 CV Preview</div>', unsafe_allow_html=True)
        
        # Store data in session state
        st.session_state.cv_data = {
            'name': name,
            'profile': profile,
            'experience': experiences,
            'education': education,
            'skills': skills,
            'activities': activities
        }
        
        # Preview information
        if name:
            st.markdown(f"**Name:** {name}")
        if profile:
            st.markdown(f"**Profile:** {profile[:100]}...")
        if any(exp.get('company') or exp.get('position') for exp in experiences):
            st.markdown(f"**Experience:** {len([exp for exp in experiences if exp.get('company') or exp.get('position')])} entries")
        if any(edu.get('degree') or edu.get('institution') for edu in education):
            st.markdown(f"**Education:** {len([edu for edu in education if edu.get('degree') or edu.get('institution')])} entries")
        if skills:
            st.markdown(f"**Skills:** {len(skills.split(','))} skills listed")
        if activities:
            st.markdown(f"**Activities:** Listed")
        
        st.markdown("---")
        
        # Generate CV button
        if st.button("🚀 Generate Professional CV", type="primary"):
            if not name or not profile or not skills:
                st.error("Please fill in all required fields (marked with *)")
            else:
                with st.spinner("Generating your professional CV..."):
                    try:
                        # Create the CV document
                        doc = create_cv_document(st.session_state.cv_data, 
                                               enhance_with_ai and groq_api_key, 
                                               groq_api_key)
                        
                        # Save to memory
                        bio = io.BytesIO()
                        doc.save(bio)
                        bio.seek(0)
                        
                        # Success message
                        st.success("✅ CV generated successfully!")
                        
                        # Download button
                        st.download_button(
                            label="📥 Download CV",
                            data=bio.getvalue(),
                            file_name=f"{name.replace(' ', '_')}_CV_{datetime.now().strftime('%Y%m%d')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                        
                        if enhance_with_ai and groq_api_key:
                            st.info("🤖 Your CV has been enhanced with AI for better professional impact!")
                        
                    except Exception as e:
                        st.error(f"Error generating CV: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("💡 **Tips for a great CV:**")
    st.markdown("• Use action verbs and quantify achievements")
    st.markdown("• Keep it concise and relevant to the job")
    st.markdown("• Proofread carefully for errors")
    st.markdown("• Tailor your CV for each application")

if __name__ == "__main__":
    main()